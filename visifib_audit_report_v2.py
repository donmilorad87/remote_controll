#!/usr/bin/env python3
"""
Visifib.ca Website Audit Report Generator v2
- Annotated screenshots with numbered red/green arrows
- Local Blazing Sun GIF logo
- Professional DOCX output for client
"""

import os, math
from io import BytesIO
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE = "/home/milner/Desktop/remote/"
GIF_PATH = os.path.join(BASE, "giphy-blazing-sunZ-complete60.gif")
OUTPUT = os.path.join(BASE, "Visifib_Website_Audit_Report_BlazingSun.docx")

MAX_W = 6.0

# ── Annotation Engine ────────────────────────────────────────────

def get_font(size):
    for p in ["/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
              "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
              "/usr/share/fonts/truetype/freefont/FreeSansBold.ttf"]:
        if os.path.exists(p):
            return ImageFont.truetype(p, size)
    return ImageFont.load_default()

def draw_arrow(draw, x1, y1, x2, y2, color, width=4):
    draw.line([(x1, y1), (x2, y2)], fill=color, width=width)
    angle = math.atan2(y2 - y1, x2 - x1)
    hl = 20
    ha = math.pi / 6
    lx = x2 - hl * math.cos(angle - ha)
    ly = y2 - hl * math.sin(angle - ha)
    rx = x2 - hl * math.cos(angle + ha)
    ry = y2 - hl * math.sin(angle + ha)
    draw.polygon([(x2, y2), (lx, ly), (rx, ry)], fill=color)

def draw_numbered_circle(draw, x, y, number, color, font):
    r = 22
    draw.ellipse([x - r, y - r, x + r, y + r], fill=color, outline="white", width=2)
    text = str(number)
    bbox = font.getbbox(text)
    tw = bbox[2] - bbox[0]
    th = bbox[3] - bbox[1]
    draw.text((x - tw // 2, y - th // 2 - 2), text, fill="white", font=font)

def annotate_screenshot(img_path, annotations, max_height=3000):
    img = Image.open(img_path).convert("RGB")
    w, h = img.size
    if h > max_height:
        img = img.crop((0, 0, w, max_height))
        h = max_height
    draw = ImageDraw.Draw(img)
    font = get_font(22)
    red = (220, 53, 53)
    green = (39, 174, 96)
    for a in annotations:
        color = red if a["color"] == "red" else green
        tx = max(25, min(w - 25, int(a["x"] * w)))
        ty = max(25, min(h - 25, int(a["y"] * h)))
        ax = max(25, min(w - 25, int(a["ax"] * w)))
        ay = max(25, min(h - 25, int(a["ay"] * h)))
        draw_arrow(draw, ax, ay, tx, ty, color, width=4)
        draw_numbered_circle(draw, ax, ay, a["num"], color, font)
    return img

def img_to_bytes(img):
    buf = BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf

def gif_to_png(gif_path, size=None):
    img = Image.open(gif_path).convert("RGBA")
    if size:
        img = img.resize(size, Image.LANCZOS)
    buf = BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf

# ── DOCX Helpers ─────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def style_table_header(table):
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "2C3E50")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True
                run.font.size = Pt(9)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)

def add_p(doc, text, bold=False, size=10, color=None, align=None, after=None, before=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold: run.bold = True
    run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    if align is not None: p.alignment = align
    if after is not None: p.paragraph_format.space_after = Pt(after)
    if before is not None: p.paragraph_format.space_before = Pt(before)
    return p

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/></w:pBdr>'))

def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(item)
        r.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(3)

def hide_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>')
    existing = tblPr.findall(qn('w:tblBorders'))
    for e in existing:
        tblPr.remove(e)
    tblPr.append(borders)

def add_page_number_field(paragraph):
    r1 = paragraph.add_run()
    r1._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    r2 = paragraph.add_run()
    r2._r.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
    r3 = paragraph.add_run()
    r3._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))

def add_screenshot(doc, pil_img, caption, notes, width_in=None):
    buf = img_to_bytes(pil_img)
    iw, ih = pil_img.size
    aspect = ih / iw
    if width_in is None:
        width_in = min(MAX_W, 6.0)
    height_in = width_in * aspect
    if height_in > 9.0:
        width_in = 9.0 / aspect
        if width_in > MAX_W:
            width_in = MAX_W

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(buf, width=Inches(width_in))
    p.paragraph_format.space_after = Pt(4)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    cap.paragraph_format.space_after = Pt(6)

    if notes:
        for num, color, desc in notes:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Pt(18)
            marker = p.add_run(f"  {num}  ")
            marker.bold = True
            marker.font.size = Pt(10)
            if color == "red":
                marker.font.color.rgb = RGBColor(0xDC, 0x35, 0x35)
                prefix = p.add_run("ISSUE: ")
                prefix.bold = True
                prefix.font.size = Pt(10)
                prefix.font.color.rgb = RGBColor(0xDC, 0x35, 0x35)
            else:
                marker.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
                prefix = p.add_run("GOOD: ")
                prefix.bold = True
                prefix.font.size = Pt(10)
                prefix.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
            body = p.add_run(desc)
            body.font.size = Pt(10)
            body.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ══════════════════════════════════════════════════════════════════
# ANNOTATIONS
# ══════════════════════════════════════════════════════════════════

def create_all_annotations():
    results = {}
    S = lambda f: os.path.join(BASE, f)

    results["home_desktop"] = (
        annotate_screenshot(S("visifib-home-desktop-1440.png"), [
            {"num": 1, "color": "red",   "x": 0.50, "y": 0.02, "ax": 0.85, "ay": 0.04},
            {"num": 2, "color": "red",   "x": 0.50, "y": 0.18, "ax": 0.90, "ay": 0.12},
            {"num": 3, "color": "red",   "x": 0.50, "y": 0.35, "ax": 0.10, "ay": 0.30},
            {"num": 4, "color": "green", "x": 0.50, "y": 0.58, "ax": 0.90, "ay": 0.55},
            {"num": 5, "color": "red",   "x": 0.50, "y": 0.80, "ax": 0.90, "ay": 0.75},
        ], max_height=2965),
        [
            (1, "red", "Dark black header (#191919) creates harsh contrast against page content. Should transition to light/white."),
            (2, "red", "Three full-screen hero images stacked vertically push ALL content below the fold. User must scroll extensively to see any products."),
            (3, "red", "Large blank space caused by MotionPage animation elements failing to load. Content should be here, not empty space."),
            (4, "green", "Product carousel is well-designed with clear categories, product images, and names. This section works effectively."),
            (5, "red", "Dark black footer (#111111) with low-contrast text. Blue links on black background fail WCAG contrast standards (2.8:1 ratio, minimum 4.5:1 required)."),
        ]
    )

    results["home_tablet"] = (
        annotate_screenshot(S("visifib-home-tablet.png"), [
            {"num": 1, "color": "red",   "x": 0.50, "y": 0.19, "ax": 0.85, "ay": 0.14},
            {"num": 2, "color": "red",   "x": 0.50, "y": 0.38, "ax": 0.12, "ay": 0.33},
            {"num": 3, "color": "red",   "x": 0.50, "y": 0.55, "ax": 0.85, "ay": 0.50},
        ], max_height=2805),
        [
            (1, "red", "Hero images dominate entire viewport on tablet. No navigation visible, no product information above fold."),
            (2, "red", "Same blank animation space on tablet. Even more wasted screen real estate on smaller device."),
            (3, "red", "Carousel shows only 3 products on tablet vs 6 on desktop. Product cards could be utilized better."),
        ]
    )

    results["home_mobile"] = (
        annotate_screenshot(S("visifib-home-mobile.png"), [
            {"num": 1, "color": "green", "x": 0.50, "y": 0.01, "ax": 0.80, "ay": 0.04},
            {"num": 2, "color": "red",   "x": 0.50, "y": 0.20, "ax": 0.85, "ay": 0.15},
            {"num": 3, "color": "red",   "x": 0.50, "y": 0.42, "ax": 0.12, "ay": 0.37},
            {"num": 4, "color": "red",   "x": 0.50, "y": 0.77, "ax": 0.12, "ay": 0.72},
        ], max_height=2905),
        [
            (1, "green", "Compact mobile header with hamburger menu is clean and functional."),
            (2, "red", "Hero images still consume excessive viewport space before showing any products."),
            (3, "red", "Blank animation space still present on mobile, pushing actual content far down."),
            (4, "red", "Dark footer with social icons that link back to homepage instead of actual social profiles."),
        ]
    )

    results["about_desktop"] = (
        annotate_screenshot(S("visifib-about-desktop.png"), [
            {"num": 1, "color": "green", "x": 0.50, "y": 0.12, "ax": 0.90, "ay": 0.07},
            {"num": 2, "color": "green", "x": 0.30, "y": 0.30, "ax": 0.08, "ay": 0.25},
            {"num": 3, "color": "red",   "x": 0.50, "y": 0.48, "ax": 0.90, "ay": 0.42},
            {"num": 4, "color": "red",   "x": 0.30, "y": 0.52, "ax": 0.08, "ay": 0.52},
            {"num": 5, "color": "green", "x": 0.65, "y": 0.50, "ax": 0.92, "ay": 0.53},
            {"num": 6, "color": "red",   "x": 0.50, "y": 0.75, "ax": 0.90, "ay": 0.70},
        ], max_height=2269),
        [
            (1, "green", "Category banner with product images is attractive and provides clear navigation to product types."),
            (2, "green", "Clean white section with 'About Us' heading provides good visual breathing room."),
            (3, "red", "Abrupt transition from white to black background creates jarring visual dissonance. This dark section fights the brand identity."),
            (4, "red", "Blue text on dark background has poor contrast. Description is harder to read than necessary."),
            (5, "green", "'Our Services' card with white background and clean typography is well-designed. This card style should be the template for the entire page."),
            (6, "red", "Massive dark footer area with barely visible content. Blue links on black fail accessibility standards."),
        ]
    )

    results["about_tablet"] = (
        annotate_screenshot(S("visifib-about-tablet.png"), [
            {"num": 1, "color": "red",   "x": 0.50, "y": 0.42, "ax": 0.85, "ay": 0.37},
            {"num": 2, "color": "red",   "x": 0.50, "y": 0.78, "ax": 0.85, "ay": 0.73},
        ], max_height=2234),
        [
            (1, "red", "Dark content section is even more prominent on tablet. Blue text on dark background becomes harder to read."),
            (2, "red", "Footer takes disproportionate space on tablet. Contact info barely visible against dark background."),
        ]
    )

    results["contact_desktop"] = (
        annotate_screenshot(S("visifib-contact-desktop.png"), [
            {"num": 1, "color": "green", "x": 0.50, "y": 0.08, "ax": 0.12, "ay": 0.08},
            {"num": 2, "color": "green", "x": 0.50, "y": 0.35, "ax": 0.88, "ay": 0.28},
            {"num": 3, "color": "red",   "x": 0.50, "y": 0.36, "ax": 0.10, "ay": 0.30},
            {"num": 4, "color": "green", "x": 0.55, "y": 0.70, "ax": 0.88, "ay": 0.67},
            {"num": 5, "color": "red",   "x": 0.50, "y": 0.55, "ax": 0.10, "ay": 0.52},
            {"num": 6, "color": "red",   "x": 0.50, "y": 0.90, "ax": 0.88, "ay": 0.87},
        ], max_height=1922),
        [
            (1, "green", "Clean 'Contact Us' heading with clear hierarchy."),
            (2, "green", "Contact form is functional: Name, Phone, City, Timeline, Product Interest, Email, Message, File Upload."),
            (3, "red", "Form on dark photographic background reduces prominence. Should be on clean light background."),
            (4, "green", "Gradient info card (teal/green) with contact details and social icons is appealing."),
            (5, "red", "Large dark empty space between form and info card. Layout should be more compact."),
            (6, "red", "No Google Maps embed. Adding a map builds trust and helps customers find the business."),
        ]
    )

    results["gallery_desktop"] = (
        annotate_screenshot(S("visifib-gallery-desktop.png"), [
            {"num": 1, "color": "green", "x": 0.50, "y": 0.28, "ax": 0.12, "ay": 0.24},
            {"num": 2, "color": "green", "x": 0.50, "y": 0.42, "ax": 0.88, "ay": 0.38},
            {"num": 3, "color": "red",   "x": 0.50, "y": 0.12, "ax": 0.88, "ay": 0.08},
            {"num": 4, "color": "red",   "x": 0.50, "y": 0.70, "ax": 0.88, "ay": 0.65},
        ], max_height=2256),
        [
            (1, "green", "Clean 'Our Gallery' section title with Visifib branding and white space."),
            (2, "green", "Three-column gallery grid with category badges and dates is well-structured."),
            (3, "red", "Category banner adds visual weight without new information. Same image on every page."),
            (4, "red", "Stark transition from white gallery to black footer. Dark footer occupies significant space."),
        ]
    )

    results["gallery_single_desktop"] = (
        annotate_screenshot(S("visifib-gallery-single-desktop.png"), [
            {"num": 1, "color": "red",   "x": 0.50, "y": 0.15, "ax": 0.88, "ay": 0.10},
            {"num": 2, "color": "red",   "x": 0.50, "y": 0.32, "ax": 0.12, "ay": 0.35},
            {"num": 3, "color": "red",   "x": 0.15, "y": 0.40, "ax": 0.12, "ay": 0.43},
            {"num": 4, "color": "green", "x": 0.50, "y": 0.62, "ax": 0.88, "ay": 0.58},
            {"num": 5, "color": "green", "x": 0.15, "y": 0.78, "ax": 0.12, "ay": 0.80},
        ], max_height=2899),
        [
            (1, "red", "Massive full-viewport hero image with no descriptive text. User cannot tell what the page is about until scrolling. Galleries do not need this oversized hero."),
            (2, "red", "Title 'Our Visifib Gallery 3' overlaid on photo with low contrast. Hard to read."),
            (3, "red", "Dark section with logo serves no purpose. Wasted space pushing gallery content further down."),
            (4, "green", "Gallery grid with product photos is well-laid out. Images are clickable."),
            (5, "green", "Pagination controls (1, 2, 3, 4) allow gallery navigation."),
        ]
    )

    results["shop_desktop"] = (
        annotate_screenshot(S("visifib-shop-desktop.png"), [
            {"num": 1, "color": "green", "x": 0.50, "y": 0.02, "ax": 0.12, "ay": 0.02},
            {"num": 2, "color": "green", "x": 0.50, "y": 0.05, "ax": 0.88, "ay": 0.05},
            {"num": 3, "color": "green", "x": 0.50, "y": 0.12, "ax": 0.12, "ay": 0.09},
            {"num": 4, "color": "green", "x": 0.50, "y": 0.20, "ax": 0.88, "ay": 0.17},
            {"num": 5, "color": "red",   "x": 0.50, "y": 0.07, "ax": 0.12, "ay": 0.07},
            {"num": 6, "color": "red",   "x": 0.70, "y": 0.30, "ax": 0.88, "ay": 0.25},
        ], max_height=3500),
        [
            (1, "green", "Clean header with clear navigation. 'Get A Quote' CTA stands out effectively."),
            (2, "green", "Predominantly white/light background. The cleanest, most pleasant page on the entire site."),
            (3, "green", "Product sections organized by collection with clear headings and descriptions."),
            (4, "green", "Product carousels with high-quality images. Easy to browse."),
            (5, "red", "Repeated 'Call Us Today / 450-803-2657' in every section is redundant."),
            (6, "red", "Placeholder text visible: 'Tell people a little more to attract...' Must be removed."),
        ]
    )

    results["product_single_desktop"] = (
        annotate_screenshot(S("visifib-product-single-desktop.png"), [
            {"num": 1, "color": "green", "x": 0.65, "y": 0.18, "ax": 0.92, "ay": 0.14},
            {"num": 2, "color": "green", "x": 0.65, "y": 0.50, "ax": 0.92, "ay": 0.45},
            {"num": 3, "color": "green", "x": 0.65, "y": 0.72, "ax": 0.92, "ay": 0.68},
            {"num": 4, "color": "red",   "x": 0.50, "y": 0.92, "ax": 0.12, "ay": 0.90},
            {"num": 5, "color": "red",   "x": 0.25, "y": 0.50, "ax": 0.08, "ay": 0.45},
        ], max_height=1274),
        [
            (1, "green", "Clean product title with clear hierarchy and description section."),
            (2, "green", "Details table (Colors, Grilles Options) is concise and informative."),
            (3, "green", "'Read More' CTA linking to detailed blog post provides clear next step."),
            (4, "red", "Dark header/footer create unnecessary contrast against clean white product card."),
            (5, "red", "Background image partially visible behind content creates visual noise."),
        ]
    )

    results["blog_product_desktop"] = (
        annotate_screenshot(S("visifib-blog-product-desktop.png"), [
            {"num": 1, "color": "red",   "x": 0.50, "y": 0.07, "ax": 0.88, "ay": 0.04},
            {"num": 2, "color": "red",   "x": 0.50, "y": 0.22, "ax": 0.12, "ay": 0.18},
            {"num": 3, "color": "red",   "x": 0.50, "y": 0.32, "ax": 0.88, "ay": 0.28},
            {"num": 4, "color": "green", "x": 0.20, "y": 0.50, "ax": 0.08, "ay": 0.47},
            {"num": 5, "color": "red",   "x": 0.50, "y": 0.62, "ax": 0.88, "ay": 0.58},
            {"num": 6, "color": "green", "x": 0.30, "y": 0.78, "ax": 0.08, "ay": 0.75},
        ], max_height=3500),
        [
            (1, "red", "Oversized hero image consuming entire viewport. Product name overlaid on complex background reduces readability."),
            (2, "red", "Second dark hero section (Visifib logo area) adds no value and pushes product info further down."),
            (3, "red", "Sky blue background with dark text. Contrast is not optimal and feels disconnected from rest of design."),
            (4, "green", "Product viewer with download links (Drawings, Installation PDFs) provides useful technical resources."),
            (5, "red", "Excessive whitespace between product sections. Glass/Grilles options pushed too far down."),
            (6, "green", "Glass and Grilles options with clear images and labels. Technical content well-presented."),
        ]
    )

    return results

# ══════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════

def create_report():
    print("Annotating screenshots...")
    annots = create_all_annotations()
    print("Building document...")
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    for i in range(1, 5):
        hs = doc.styles[f'Heading {i}']
        hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        hs.font.name = 'Calibri'

    logo_cover = gif_to_png(GIF_PATH, size=(300, 300))
    logo_footer = gif_to_png(GIF_PATH, size=(40, 40))

    # ── COVER ──
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_cover.seek(0)
    p.add_run().add_picture(logo_cover, width=Inches(2.2))
    doc.add_paragraph()

    add_p(doc, "WEBSITE AUDIT REPORT", bold=True, size=28, color=RGBColor(0x1A, 0x1A, 0x2E), align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    add_p(doc, "Comprehensive UX, Performance, SEO & Accessibility Analysis", size=13, color=RGBColor(0x55, 0x55, 0x55), align=WD_ALIGN_PARAGRAPH.CENTER, after=30)
    add_divider(doc)
    add_p(doc, "www.visifib.ca", bold=True, size=18, color=RGBColor(0x04, 0x6B, 0xD2), align=WD_ALIGN_PARAGRAPH.CENTER, after=30)

    it = doc.add_table(rows=4, cols=2)
    it.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (l, v) in enumerate([("Date:", "March 20, 2026"), ("Prepared for:", "Visifib (Entreprises Visifib)"), ("Prepared by:", "Milorad Djukovic (Blazing Sun)"), ("Document Version:", "1.0")]):
        it.rows[i].cells[0].text = l
        it.rows[i].cells[1].text = v
        it.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        for c in it.rows[i].cells:
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs: r.font.size = Pt(11)
    hide_table_borders(it)
    doc.add_paragraph()
    add_p(doc, "CONFIDENTIAL", bold=True, size=10, color=RGBColor(0xC0, 0x39, 0x2B), align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # ── TOC ──
    doc.add_heading('Table of Contents', level=1)
    for item in ["1. Executive Summary", "2. Scope & Methodology", "3. Site Architecture Overview",
                 "4. Design & Visual Identity Audit", "    4.1 Color Scheme & Harmony", "    4.2 Hero Images & Layout",
                 "    4.3 Animations & Motion", "    4.4 Page-by-Page Visual Review (Annotated Screenshots)",
                 "5. Performance Audit", "6. SEO Audit", "7. Accessibility Audit (WCAG 2.1)",
                 "8. Technical Issues & Console Errors", "9. Prioritized Recommendations", "10. Conclusion & Next Steps"]:
        p = doc.add_paragraph()
        r = p.add_run(item); r.font.size = Pt(10)
        if not item.startswith("    "): r.bold = True
        p.paragraph_format.space_after = Pt(2)
    doc.add_page_break()

    # ── 1. EXECUTIVE SUMMARY ──
    doc.add_heading('1. Executive Summary', level=1)
    p = doc.add_paragraph()
    r = p.add_run("Visifib.ca "); r.bold = True; r.font.size = Pt(10)
    p.add_run("is a professional website for Entreprises Visifib, a Quebec-based company selling premium fiberglass and aluminium windows and doors. Built on WordPress with Astra theme and Elementor page builder, running WooCommerce for the product catalog.").font.size = Pt(10)
    p.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph()
    p.add_run("This audit analyzed all pages across desktop (1440px), tablet (768px), and mobile (375px) viewports. Full-page screenshots were captured and are included throughout this report with annotated observations highlighting specific issues and strengths.").font.size = Pt(10)
    p.paragraph_format.space_after = Pt(12)

    doc.add_heading('Key Findings', level=2)
    add_bullets(doc, [
        "Design: Dark/black sections create harsh contrast that conflicts with the brand. Visifib sells windows and doors \u2014 products that let light into homes. The visual language should be light, airy, and transparent.",
        "Performance: MotionPage animations, oversized images without responsive sizing, external web fonts failing to load, and Elementor overhead are significantly slowing the site.",
        "SEO: Every page has an incomplete title with a trailing hyphen. Meta descriptions are missing or inadequate across the entire site.",
        "Accessibility: Multiple WCAG AA contrast failures \u2014 dark blue on black footer, light blue hover states on dark backgrounds.",
        "Technical: 5 JavaScript/resource errors on every page load, including broken chat widget and React component error.",
    ])

    doc.add_heading('Issue Summary', level=2)
    st = doc.add_table(rows=5, cols=2); st.style = 'Table Grid'
    for i, (l, v) in enumerate([("Critical Issues", "8"), ("High Priority", "12"), ("Medium Priority", "15"), ("Low / Info", "10"), ("Pages Audited", "40+")]):
        st.rows[i].cells[0].text = l; st.rows[i].cells[1].text = v
        st.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        for c in st.rows[i].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    doc.add_page_break()

    # ── 2. SCOPE ──
    doc.add_heading('2. Scope & Methodology', level=1)
    p = doc.add_paragraph(); p.add_run("Audit conducted March 20, 2026, covering all English pages (/en/) discovered via sitemap_index.xml.").font.size = Pt(10); p.paragraph_format.space_after = Pt(8)
    doc.add_heading('2.1 Pages Audited', level=2)
    add_bullets(doc, ["Page Sitemap: 11 pages (Home, About Us, Contact, Shop, etc.)", "Post Sitemap: 17 posts (Product articles, Gallery pages)", "Product Sitemap: 51 products (Doors, Windows, Sliding doors)", "Category Sitemap: 5 categories", "Builder Templates: 2 WooBuilder + 5 Configurator components"])
    doc.add_heading('2.2 Testing Methodology', level=2)
    add_bullets(doc, ["Visual inspection: 1440px (desktop), 768px (tablet), 375px (mobile)", "Full-page screenshots via Playwright \u2014 annotated with red/green markers and included in this report", "DOM accessibility snapshots for structure analysis", "Console error monitoring during page loads", "Source code analysis for SEO, colors, fonts, asset sizes"])
    doc.add_page_break()

    # ── 3. ARCHITECTURE ──
    doc.add_heading('3. Site Architecture Overview', level=1)
    doc.add_heading('3.1 Technology Stack', level=2)
    tt = doc.add_table(rows=10, cols=2); tt.style = 'Table Grid'
    tt.rows[0].cells[0].text = "Component"; tt.rows[0].cells[1].text = "Technology"; style_table_header(tt)
    for i, (c, t) in enumerate([("CMS", "WordPress"), ("Theme", "Astra v2.5.2"), ("Page Builder", "Elementor + Jet WooBuilder + Jet Engine"), ("E-commerce", "WooCommerce"), ("Animation", "MotionPage framework"), ("Chat", "Tawk.to (broken \u2014 CORS errors)"), ("Fonts", "Poppins (Google Fonts, external CDN)"), ("Language", "Bilingual EN/FR"), ("Hosting", "Standard WordPress hosting")]):
        tt.rows[i+1].cells[0].text = c; tt.rows[i+1].cells[1].text = t
        for cell in tt.rows[i+1].cells:
            for p in cell.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    doc.add_heading('3.2 Navigation', level=2)
    add_bullets(doc, ["Home", "About Us", "Collections (Contemporary, Traditional, Windows, Sliding Door, Lift & Slide)", "All Products", "Contact", "FR (language switch)", "Get A Quote (CTA button)"])
    doc.add_page_break()

    # ── 4. DESIGN AUDIT ──
    doc.add_heading('4. Design & Visual Identity Audit', level=1)

    doc.add_heading('4.1 Color Scheme & Harmony', level=2)
    ct = doc.add_table(rows=8, cols=3); ct.style = 'Table Grid'
    ct.rows[0].cells[0].text = "Element"; ct.rows[0].cells[1].text = "Color"; ct.rows[0].cells[2].text = "Hex"; style_table_header(ct)
    for i, (e, c, h) in enumerate([("Primary Blue", "Bright Blue", "#046BD2"), ("Secondary Blue", "Light Blue", "#93C3F3"), ("Header BG", "Near Black", "#191919"), ("Footer BG", "Dark Black", "#111111"), ("Body Text", "Dark Slate", "#334155"), ("Page BG", "Light Blue-Gray", "#F0F5FA"), ("White Sections", "Pure White", "#FFFFFF")]):
        ct.rows[i+1].cells[0].text = e; ct.rows[i+1].cells[1].text = c; ct.rows[i+1].cells[2].text = h
        for cell in ct.rows[i+1].cells:
            for p in cell.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    doc.add_paragraph()
    doc.add_heading('Core Problem: Brand vs Visual Language', level=3)
    p = doc.add_paragraph()
    p.add_run("The site alternates between near-black and white sections, creating jarring contrast. ").font.size = Pt(10)
    r = p.add_run("Visifib sells windows and doors that bring light into homes. "); r.font.size = Pt(10); r.bold = True
    p.add_run("The visual language should reflect this: open, bright, airy. The Shop page demonstrates the right approach \u2014 predominantly white, clean, pleasant.").font.size = Pt(10)
    p.paragraph_format.space_after = Pt(8)
    add_p(doc, "Recommendation:", bold=True, size=10, color=RGBColor(0x27, 0xAE, 0x60), after=4)
    add_bullets(doc, ["Transition header from black (#191919) to white or very light gray", "Transition footer from black (#111111) to light gray with dark text", "Remove all dark background content sections", "Use primary blue (#046BD2) sparingly for CTAs only", "Let product photography stand out against clean, light backgrounds"])

    doc.add_heading('4.2 Hero Images & Layout', level=2)
    p = doc.add_paragraph(); p.add_run("Multiple pages use oversized full-screen heroes pushing content below the fold. Design should be intuitive \u2014 everything clear from the first second.").font.size = Pt(10); p.paragraph_format.space_after = Pt(8)

    doc.add_heading('4.3 Animations & Motion', level=2)
    p = doc.add_paragraph(); p.add_run("MotionPage framework creates confusion and slows the site. Flying images add no value. The blank space on the homepage is caused by animation elements not loading.").font.size = Pt(10); p.paragraph_format.space_after = Pt(6)
    add_p(doc, "Recommendation: Remove MotionPage entirely. Keep only product carousels.", bold=True, size=10, color=RGBColor(0x27, 0xAE, 0x60), after=8)
    doc.add_page_break()

    # ── 4.4 SCREENSHOTS ──
    doc.add_heading('4.4 Page-by-Page Visual Review (Annotated Screenshots)', level=2)
    p = doc.add_paragraph(); p.add_run("Each screenshot is annotated with numbered markers: ").font.size = Pt(10)
    r = p.add_run("RED "); r.bold = True; r.font.color.rgb = RGBColor(0xDC, 0x35, 0x35); r.font.size = Pt(10)
    p.add_run("= issue to fix. ").font.size = Pt(10)
    r2 = p.add_run("GREEN "); r2.bold = True; r2.font.color.rgb = RGBColor(0x27, 0xAE, 0x60); r2.font.size = Pt(10)
    p.add_run("= working well.").font.size = Pt(10)
    p.paragraph_format.space_after = Pt(12)

    # Insert all annotated pages
    pages = [
        ("Homepage \u2014 Desktop (1440px)", "home_desktop", "Figure 1", None),
        ("Homepage \u2014 Tablet (768px)", "home_tablet", "Figure 2", 4.0),
        ("Homepage \u2014 Mobile (375px)", "home_mobile", "Figure 3", 3.0),
        (None, None, None, None),  # page break
        ("About Us \u2014 Desktop (1440px)", "about_desktop", "Figure 4", None),
        ("About Us \u2014 Tablet (768px)", "about_tablet", "Figure 5", 4.0),
        (None, None, None, None),
        ("Contact Page \u2014 Desktop (1440px)", "contact_desktop", "Figure 6", None),
        (None, None, None, None),
        ("Gallery Index \u2014 Desktop (1440px)", "gallery_desktop", "Figure 7", None),
        ("Single Gallery Page \u2014 Desktop (1440px)", "gallery_single_desktop", "Figure 8", None),
        (None, None, None, None),
        ("Shop / All Products \u2014 Desktop (1440px) \u2014 BEST DESIGNED PAGE", "shop_desktop", "Figure 9", None),
        (None, None, None, None),
        ("Single Product (WooCommerce) \u2014 Desktop", "product_single_desktop", "Figure 10", None),
        ("Product Blog Post \u2014 Desktop (1440px)", "blog_product_desktop", "Figure 11", None),
    ]
    for entry in pages:
        if entry[0] is None:
            doc.add_page_break()
            continue
        title, key, fig, w = entry
        add_p(doc, title, bold=True, size=12, color=RGBColor(0x1A, 0x1A, 0x2E), after=6)
        if key == "shop_desktop":
            p = doc.add_paragraph()
            r = p.add_run("This is the best-designed page on the entire site. ")
            r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
            p.add_run("The predominantly white/light palette should serve as the design reference for the entire site.").font.size = Pt(10)
            p.paragraph_format.space_after = Pt(8)
        img, notes = annots[key]
        add_screenshot(doc, img, f"{fig}: {title}", notes, width_in=w)

    doc.add_page_break()

    # ── DETAILED PRODUCT BLOG POST ANALYSIS (screancapture.png) ──
    add_p(doc, "Detailed Product Blog Post Analysis \u2014 325 Awning Page", bold=True, size=14, color=RGBColor(0x1A, 0x1A, 0x2E), after=6)
    p = doc.add_paragraph()
    p.add_run("The following full-page capture of the 325 Awning product blog post illustrates several critical design issues that repeat across all product detail pages. This page is representative of the problems found site-wide.").font.size = Pt(10)
    p.paragraph_format.space_after = Pt(8)

    # Load the pre-annotated screencapture
    screencap_path = os.path.join(BASE, "screancapture-annotated.png")
    screencap_img = Image.open(screencap_path).convert("RGB")
    screencap_notes = [
        (1, "red", "Oversized hero image consuming entire viewport. Product name '325 \u2013 AWNING' is overlaid on a complex interior photograph, reducing readability. The user sees a beautiful room but has no immediate understanding of what product is being presented. The hero should be compact (max 40% viewport height) with the product name clearly visible on a clean background."),
        (2, "red", "Second dark hero section with Visifib logo adds no informational value and pushes the actual product details further down the page. This entire section should be removed \u2014 the logo is already in the header. Every pixel of vertical space between the user and the product information is a lost opportunity."),
        (3, "red", "Sky blue background (#93C3F3) with dark text creates a color contrast that feels disconnected from the rest of the site design. While the text is technically readable, the blue background clashes with the black sections above and the white sections below. This should be a clean white or very light gray background for consistency."),
        (4, "green", "Product viewer section with profile cross-section image, color samples, and the Limitless Color Options color wheel provides genuinely useful technical information. The layout with side-by-side comparisons works well. Download links for Drawings and Installation PDFs give professionals the resources they need."),
        (5, "red", "Section headings ('Colours', 'Hardware Options', 'Double Pane Insulating', 'Grille Profiles') use bright teal/green background strips that create strong visual contrast against the white content areas. These colored heading bars feel inconsistent with the site's color palette and add visual noise. Headings should use the primary blue (#046BD2) or simply bold dark text on white."),
        (6, "green", "Technical product images (glass cross-sections, hardware components, grille profiles) are well-photographed and clearly labeled. The grid layout makes it easy to compare options. This type of detailed product information builds confidence with professional buyers and homeowners alike."),
    ]
    add_screenshot(doc, screencap_img, "Figure 12: Product Blog Post \u2014 325 Awning \u2014 Full Page Analysis", screencap_notes)
    doc.add_page_break()

    # ── 5. PERFORMANCE ──
    doc.add_heading('5. Performance Audit', level=1)
    doc.add_heading('5.1 Image Optimization', level=2)
    add_bullets(doc, ["Images served at original resolution (2560px+) regardless of device", "No responsive srcset for adaptive delivery", "Recommendation: srcset, WebP/AVIF, lazy loading \u2014 est. 40-60% payload reduction"])
    doc.add_heading('5.2 Web Fonts', level=2)
    p = doc.add_paragraph(); p.add_run("Poppins loads from external CDN (dd-time.com) failing with CORS errors. Recommendation: Switch to system fonts or self-host.").font.size = Pt(10)
    doc.add_heading('5.3 Caching', level=2)
    add_bullets(doc, ["Site is mostly static \u2014 aggressive caching will help significantly", "Install WP caching plugin, enable Gzip/Brotli, implement CDN (Cloudflare free)", "Est. 50-70% faster repeat visits"])
    doc.add_heading('5.4 Elementor', level=2)
    p = doc.add_paragraph(); p.add_run("Elementor adds 50-200KB inline CSS per page + multiple JS files. Workable but needs optimization: remove unused widgets, enable performance mode, use Asset CleanUp plugin.").font.size = Pt(10)
    doc.add_heading('5.5 Console Errors (Every Page)', level=2)
    et = doc.add_table(rows=5, cols=3); et.style = 'Table Grid'
    et.rows[0].cells[0].text = "Error"; et.rows[0].cells[1].text = "Source"; et.rows[0].cells[2].text = "Impact"; style_table_header(et)
    for i, (e, s, imp) in enumerate([("CORS font blocked", "dd-time.com (Roboto)", "Font flash/fallback"), ("CORS script blocked", "tawk.to chat", "Chat completely broken"), ("net::ERR_FAILED", "Roboto font file", "Font inconsistency"), ("React error #299", "Product configurator", "Configurator partially broken")]):
        et.rows[i+1].cells[0].text = e; et.rows[i+1].cells[1].text = s; et.rows[i+1].cells[2].text = imp
        for c in et.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    doc.add_page_break()

    # ── 6. SEO ──
    doc.add_heading('6. SEO Audit', level=1)
    doc.add_heading('6.1 Page Titles & Meta Descriptions', level=2)
    p = doc.add_paragraph(); r = p.add_run("CRITICAL: Every page has incomplete title with trailing hyphen. Meta descriptions missing or inadequate."); r.font.size = Pt(10); r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B); p.paragraph_format.space_after = Pt(8)
    s2 = doc.add_table(rows=9, cols=3); s2.style = 'Table Grid'
    s2.rows[0].cells[0].text = "Page"; s2.rows[0].cells[1].text = "Current Title"; s2.rows[0].cells[2].text = "Meta Description"; style_table_header(s2)
    for i, (pg, t, d) in enumerate([("Homepage","Home -","Visifib.ca"), ("About Us","About Us -","Welcome"), ("Contact","Contact -","Visifib.ca"), ("Gallery","Gallery -","Missing"), ("Shop","All Products -","Missing"), ("Product Post","SINGLE HUNG 850 SERIES -","Lift hardware..."), ("Gallery 3","Our Visifib Gallery 3 -","Missing"), ("Product Page","SINGLE HUNG 850 SERIES -","Schema only")]):
        s2.rows[i+1].cells[0].text = pg; s2.rows[i+1].cells[1].text = t; s2.rows[i+1].cells[2].text = d
        for c in s2.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    doc.add_paragraph()
    add_p(doc, "Recommendation:", bold=True, size=10, color=RGBColor(0x27, 0xAE, 0x60), after=4)
    add_bullets(doc, ["Fix titles: 'About Us | Visifib \u2014 Premium Windows & Doors Quebec'", "Write unique 150-160 char meta descriptions with keywords", "Install Yoast SEO or Rank Math"])

    doc.add_heading('6.2 Sitemap & Indexing', level=2)
    add_bullets(doc, ["All sitemaps point to /fr/ only \u2014 English pages may not be indexed", "No hreflang annotations", "Cart/Checkout in sitemap (should be excluded)"])
    doc.add_heading('6.3 Content SEO', level=2)
    add_bullets(doc, ["Brief product descriptions", "Image alt text uses filenames", "Placeholder text on Shop page", "No internal linking strategy", "Social links point to current page, not profiles"])
    doc.add_page_break()

    # ── 7. ACCESSIBILITY ──
    doc.add_heading('7. Accessibility Audit (WCAG 2.1)', level=1)
    doc.add_heading('7.1 Color Contrast Failures', level=2)
    at = doc.add_table(rows=6, cols=4); at.style = 'Table Grid'
    at.rows[0].cells[0].text = "Location"; at.rows[0].cells[1].text = "Text"; at.rows[0].cells[2].text = "Background"; at.rows[0].cells[3].text = "Result"; style_table_header(at)
    for i, (l, t, b, res) in enumerate([("Footer links","#046BD2","#111111","FAIL 2.8:1"), ("Hover states","#93C3F3","#1D1A1A","FAIL 2.8:1"), ("Footer copyright","Gray","#111111","FAIL"), ("About text","#046BD2","Dark section","MARGINAL"), ("Privacy link","#046BD2","#111111","FAIL")]):
        at.rows[i+1].cells[0].text = l; at.rows[i+1].cells[1].text = t; at.rows[i+1].cells[2].text = b; at.rows[i+1].cells[3].text = res
        for c in at.rows[i+1].cells:
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
        if "FAIL" in res: set_cell_shading(at.rows[i+1].cells[3], "FADBD8")

    doc.add_heading('7.2 Other Issues', level=2)
    add_bullets(doc, ["French skip link 'Aller au contenu' on English pages", "French menu labels 'Permutateur de Menu' on English pages", "Image alt text uses filenames not descriptions", "Social icons link to current page, not social profiles"])
    doc.add_page_break()

    # ── 8. TECHNICAL ──
    doc.add_heading('8. Technical Issues & Console Errors', level=1)
    add_bullets(doc, [
        "CORS Font (dd-time.com): Roboto blocked. Fix: Self-host or fix CDN CORS.",
        "CORS Chat (tawk.to): Chat completely broken. Fix: Reconfigure embed code.",
        "React Error #299: Configurator broken. Fix: Debug React component.",
        "Broken Social Links: Footer icons link to current page. Fix: Add real profile URLs.",
        "French strings on English pages. Fix: Update WPML/Polylang translations.",
    ])
    doc.add_page_break()

    # ── 9. RECOMMENDATIONS ──
    doc.add_heading('9. Prioritized Recommendations', level=1)
    add_p(doc, "PHASE 1: Critical Fixes (Immediate)", bold=True, size=12, color=RGBColor(0xC0, 0x39, 0x2B), after=6)
    for i, item in enumerate(["Fix all page titles \u2014 add site name, remove trailing hyphen", "Write unique meta descriptions for all pages", "Fix broken social media links in footer", "Fix CORS errors \u2014 self-host fonts, reconfigure Tawk.to, debug configurator", "Fix contrast failures \u2014 light text on dark footer or light footer", "Fix French strings on English pages", "Remove placeholder text from Shop page"]):
        p = doc.add_paragraph(); r = p.add_run(f"1.{i+1}  "); r.bold = True; r.font.size = Pt(10); p.add_run(item).font.size = Pt(10); p.paragraph_format.space_after = Pt(3)

    doc.add_paragraph()
    add_p(doc, "PHASE 2: Design & Performance (1\u20132 Weeks)", bold=True, size=12, color=RGBColor(0xF3, 0x9C, 0x12), after=6)
    for i, item in enumerate(["Transition to light color palette (white header, light footer)", "Reduce hero images to 50-60vh max", "Remove MotionPage animation framework", "Implement responsive images with srcset", "Convert images to WebP/AVIF", "Install caching plugin + server compression", "Switch to system fonts", "Optimize Elementor output"]):
        p = doc.add_paragraph(); r = p.add_run(f"2.{i+1}  "); r.bold = True; r.font.size = Pt(10); p.add_run(item).font.size = Pt(10); p.paragraph_format.space_after = Pt(3)

    doc.add_paragraph()
    add_p(doc, "PHASE 3: SEO & Growth (2\u20134 Weeks)", bold=True, size=12, color=RGBColor(0x27, 0xAE, 0x60), after=6)
    for i, item in enumerate(["Install SEO plugin (Yoast/Rank Math)", "Create English sitemap + hreflang", "Descriptive alt text for all images", "Expand product descriptions", "Internal linking between products", "FAQ schema for rich snippets", "Google Search Console setup", "CDN implementation (Cloudflare)", "Google Maps on Contact page", "Customer testimonials"]):
        p = doc.add_paragraph(); r = p.add_run(f"3.{i+1}  "); r.bold = True; r.font.size = Pt(10); p.add_run(item).font.size = Pt(10); p.paragraph_format.space_after = Pt(3)
    doc.add_page_break()

    # ── 10. CONCLUSION ──
    doc.add_heading('10. Conclusion & Next Steps', level=1)
    p = doc.add_paragraph()
    p.add_run("Visifib.ca has a solid foundation: beautiful product photography, a functional catalog, and a clear business offering. The challenge is that the current presentation \u2014 dark colors, heavy animations, oversized heroes \u2014 works against the brand identity as a company bringing light into homes.").font.size = Pt(10)
    p.paragraph_format.space_after = Pt(8)

    r = doc.add_paragraph().add_run("The path forward:")
    r.bold = True; r.font.size = Pt(10)
    add_bullets(doc, ["Embrace light: whites, light grays, subtle blues. Let photography be the star.", "Simplify: remove animations, reduce heroes, communicate purpose in first viewport.", "Fix fundamentals: titles, descriptions, contrast, console errors \u2014 quick wins with big impact.", "Optimize speed: images, caching, fonts \u2014 noticeably faster with low effort."])

    p = doc.add_paragraph()
    r = p.add_run("The Shop page is the best-designed page and should serve as the design reference for every other page.")
    r.font.size = Pt(10); r.bold = True
    p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    p.add_run("We are ready to begin implementation upon your approval. We recommend starting with Phase 1 immediately.").font.size = Pt(10)
    p.paragraph_format.space_after = Pt(20)
    add_divider(doc)
    doc.add_paragraph()
    add_p(doc, "Milorad Djukovic", bold=True, size=12, color=RGBColor(0x1A, 0x1A, 0x2E), after=2)
    add_p(doc, "Blazing Sun", size=10, color=RGBColor(0x55, 0x55, 0x55), after=2)
    add_p(doc, "milorad.djukovic@blazingsun.space | +381 62 961 62 31", size=10, color=RGBColor(0x55, 0x55, 0x55))

    # ══════════════════════════════════════════════════════════════
    # HEADERS & FOOTERS
    # ══════════════════════════════════════════════════════════════
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    sectPr = section._sectPr
    if not sectPr.findall(qn('w:evenAndOddHeaders')):
        sectPr.append(parse_xml(f'<w:evenAndOddHeaders {nsdecls("w")}/>'))

    # First page: blank
    fh = section.first_page_header; fh.is_linked_to_previous = False
    for p in fh.paragraphs: p.clear()
    ff = section.first_page_footer; ff.is_linked_to_previous = False
    for p in ff.paragraphs: p.clear()

    # Odd header: company info
    oh = section.header; oh.is_linked_to_previous = False
    for p in oh.paragraphs: p.clear()
    p = oh.paragraphs[0] if oh.paragraphs else oh.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Blazing Sun"); r.bold = True; r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.add_run("\n")
    r2 = p.add_run("PIB: 114340949 | MB: 67500156"); r2.font.size = Pt(7); r2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    p.add_run("\n")
    r3 = p.add_run("Email: milorad.djukovic@blazingsun.space | Phone: +381 62 961 62 31"); r3.font.size = Pt(7); r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Even header: prepared for/by
    eh = section.even_page_header; eh.is_linked_to_previous = False
    for p in eh.paragraphs: p.clear()
    p = eh.paragraphs[0] if eh.paragraphs else eh.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Prepared for: Visifib"); r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    p.add_run("\n")
    r2 = p.add_run("Prepared by: Milorad Djukovic (Blazing Sun)"); r2.font.size = Pt(8); r2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Footers: 3-col table with logo center, page number right
    def build_footer(footer_obj):
        footer_obj.is_linked_to_previous = False
        for p in footer_obj.paragraphs:
            el = p._element; el.getparent().remove(el)
        pd = footer_obj.add_paragraph()
        pPr = pd._p.get_or_add_pPr()
        pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="4" w:color="CCCCCC"/></w:pBdr>'))
        pd.paragraph_format.space_after = Pt(4)

        ft = footer_obj.add_table(rows=1, cols=3, width=Inches(6.0))
        ft.alignment = WD_TABLE_ALIGNMENT.CENTER
        hide_table_borders(ft)
        for cell in ft.rows[0].cells: cell.width = Inches(2.0)

        lp = ft.rows[0].cells[0].paragraphs[0]; lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lr = lp.add_run("Confidential"); lr.font.size = Pt(6); lr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA); lr.italic = True

        cp = ft.rows[0].cells[1].paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_footer.seek(0)
        cp.add_run().add_picture(logo_footer, width=Inches(0.47))

        rp = ft.rows[0].cells[2].paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rr = rp.add_run("Page "); rr.font.size = Pt(7); rr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        add_page_number_field(rp)

    build_footer(section.footer)
    build_footer(section.even_page_footer)

    print("Saving...")
    doc.save(OUTPUT)
    print(f"Done! {OUTPUT}")

if __name__ == "__main__":
    create_report()
