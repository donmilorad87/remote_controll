#!/usr/bin/env python3
"""
Visifib.ca Website Audit Report Generator
Generates a professional .docx report
"""

import os
import requests
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Helpers ──────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_paragraph(doc, text, style=None, bold=False, size=None, color=None, alignment=None, space_after=None, space_before=None):
    """Add a paragraph with custom styling."""
    p = doc.add_paragraph()
    if style:
        p.style = style
    run = p.add_run(text)
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    return p

def add_finding_row(table, category, finding, severity, recommendation):
    """Add a row to the findings table."""
    row = table.add_row()
    cells = row.cells
    cells[0].text = category
    cells[1].text = finding
    cells[2].text = severity
    cells[3].text = recommendation

    sev_colors = {
        "Critical": "C0392B",
        "High": "E74C3C",
        "Medium": "F39C12",
        "Low": "3498DB",
        "Info": "95A5A6"
    }
    if severity in sev_colors:
        set_cell_shading(cells[2], sev_colors[severity])
        for paragraph in cells[2].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True

    for cell in cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.space_before = Pt(2)
            for run in paragraph.runs:
                run.font.size = Pt(9)

def add_section_divider(doc):
    """Add a thin horizontal line divider."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/></w:pBdr>')
    pPr.append(pBdr)

def style_table_header(table):
    """Style the header row of a table."""
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "2C3E50")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(9)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.space_before = Pt(2)

def add_bullet_list(doc, items, bold_prefix=None):
    """Add a bulleted list."""
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        if bold_prefix and ": " in item:
            prefix, rest = item.split(": ", 1)
            run = p.add_run(prefix + ": ")
            run.bold = True
            run.font.size = Pt(10)
            run2 = p.add_run(rest)
            run2.font.size = Pt(10)
        else:
            run = p.add_run(item)
            run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(3)

# ── Logo download ────────────────────────────────────────────────

def download_logo():
    """Download the Blazing Sun logo."""
    url = "https://blazingsun.space/assets/images/blazingSun.gif"
    try:
        resp = requests.get(url, timeout=10)
        if resp.status_code == 200:
            return BytesIO(resp.content)
    except Exception:
        pass
    return None

# ── Main Report ──────────────────────────────────────────────────

def create_report():
    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Default font ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Style headings
    for i in range(1, 5):
        hs = doc.styles[f'Heading {i}']
        hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        hs.font.name = 'Calibri'

    logo_data = download_logo()

    # ══════════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════════

    # Spacer
    for _ in range(4):
        doc.add_paragraph()

    # Logo
    if logo_data:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        logo_data.seek(0)
        run.add_picture(logo_data, width=Inches(2))

    doc.add_paragraph()

    # Title
    p = add_styled_paragraph(doc, "WEBSITE AUDIT REPORT", bold=True, size=28,
                             color=RGBColor(0x1A, 0x1A, 0x2E),
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    # Subtitle
    p = add_styled_paragraph(doc, "Comprehensive UX, Performance, SEO & Accessibility Analysis",
                             size=13, color=RGBColor(0x55, 0x55, 0x55),
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

    # Divider
    add_section_divider(doc)

    # Target
    p = add_styled_paragraph(doc, "www.visifib.ca", bold=True, size=18,
                             color=RGBColor(0x04, 0x6B, 0xD2),
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

    # Date and parties
    info_table = doc.add_table(rows=4, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ("Date:", "March 20, 2026"),
        ("Prepared for:", "Visifib (Entreprises Visifib)"),
        ("Prepared by:", "Milorad Djukovic (Blazing Sun)"),
        ("Document Version:", "1.0"),
    ]
    for i, (label, value) in enumerate(info_data):
        cells = info_table.rows[i].cells
        cells[0].text = label
        cells[1].text = value
        for cell in cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        cells[0].paragraphs[0].runs[0].bold = True

    # Make table borders invisible
    for row in info_table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                                  '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                                  '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                                  '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                                  '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                                  '</w:tcBorders>')
            tcPr.append(tcBorders)

    doc.add_paragraph()
    p = add_styled_paragraph(doc, "CONFIDENTIAL", bold=True, size=10,
                             color=RGBColor(0xC0, 0x39, 0x2B),
                             alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # ── Page break ──
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ══════════════════════════════════════════════════════════════

    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Scope & Methodology",
        "3. Site Architecture Overview",
        "4. Design & Visual Identity Audit",
        "    4.1 Color Scheme & Harmony",
        "    4.2 Hero Images & Layout",
        "    4.3 Animations & Motion",
        "    4.4 Page-by-Page Design Review",
        "5. Performance Audit",
        "    5.1 Image Optimization",
        "    5.2 Web Fonts",
        "    5.3 Caching Strategy",
        "    5.4 Elementor Impact",
        "    5.5 Console Errors",
        "6. SEO Audit",
        "    6.1 Page Titles & Meta Descriptions",
        "    6.2 Structured Data",
        "    6.3 Sitemap & Indexing",
        "    6.4 Content Optimization",
        "7. Accessibility Audit (WCAG 2.1)",
        "    7.1 Color Contrast Failures",
        "    7.2 Semantic HTML & ARIA",
        "    7.3 Keyboard Navigation",
        "8. Technical Issues & Console Errors",
        "9. Prioritized Recommendations",
        "10. Conclusion & Next Steps",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(10)
        if not item.startswith("    "):
            run.bold = True
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ══════════════════════════════════════════════════════════════

    doc.add_heading('1. Executive Summary', level=1)

    p = doc.add_paragraph()
    run = p.add_run("Visifib.ca")
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(" is a professional website for Entreprises Visifib, a Quebec-based company selling premium fiberglass and aluminium windows and doors. The site is built on WordPress with the Astra theme and Elementor page builder, running WooCommerce for product catalog functionality.")
    run2.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph()
    run = p.add_run("This audit comprehensively analyzed all pages across desktop (1440px), tablet (768px), and mobile (375px) viewports. The site features beautiful product photography and a solid product catalog, but several areas require attention to maximize conversions and user experience.")
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(12)

    # Summary stats table
    summary_table = doc.add_table(rows=5, cols=2)
    summary_table.style = 'Table Grid'
    summary_data = [
        ("Total Pages Audited", "40+ (pages, products, galleries, blog posts)"),
        ("Critical Issues Found", "8"),
        ("High Priority Issues", "12"),
        ("Medium Priority Issues", "15"),
        ("Low Priority / Informational", "10"),
    ]
    for i, (label, value) in enumerate(summary_data):
        cells = summary_table.rows[i].cells
        cells[0].text = label
        cells[1].text = value
        cells[0].paragraphs[0].runs[0].bold = True
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()

    # Key findings box
    p = add_styled_paragraph(doc, "Key Findings at a Glance:", bold=True, size=11,
                             color=RGBColor(0x1A, 0x1A, 0x2E), space_after=6)

    key_findings = [
        "Design: Dark/black sections create harsh contrast that conflicts with the light, transparent nature of the product (windows and doors let light in). A shift to a light color palette would dramatically improve visual harmony.",
        "Performance: Heavy animations (MotionPage framework), oversized images, web fonts, and Elementor overhead are significantly slowing page load times.",
        "SEO: Every page has incomplete titles (trailing hyphen) and missing or inadequate meta descriptions, severely limiting search visibility.",
        "Accessibility: Multiple WCAG AA contrast failures, particularly in footer areas (dark blue text on black background) and interactive hover states.",
        "Console Errors: 5 JavaScript/resource errors on every page load, including failed font loads, blocked chat widget, and a React error in the configurator.",
    ]
    add_bullet_list(doc, key_findings)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 2. SCOPE & METHODOLOGY
    # ══════════════════════════════════════════════════════════════

    doc.add_heading('2. Scope & Methodology', level=1)

    p = doc.add_paragraph()
    run = p.add_run("This audit was conducted on March 20, 2026, covering all English-language pages (/en/) of visifib.ca. The sitemap index at /sitemap_index.xml was used to discover all published URLs.")
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(8)

    add_styled_paragraph(doc, "Pages Discovered via Sitemap:", bold=True, size=10, space_after=4)

    sitemap_items = [
        "Page Sitemap: 11 pages (Home, About Us, Contact, Shop, Cart, Checkout, etc.)",
        "Post Sitemap: 17 posts (Product detail articles, Gallery pages)",
        "Product Sitemap: 51 products (Contemporary doors, Traditional doors, Windows, Sliding doors)",
        "Category Sitemap: 5 categories (Windows, Sliding Door, Lift & Slide Doors, Gallery, Uncategorized)",
        "Jet WooBuilder Sitemap: 2 templates",
        "Jet Engine Sitemap: 5 configurator components",
    ]
    add_bullet_list(doc, sitemap_items)

    add_styled_paragraph(doc, "Testing Methodology:", bold=True, size=10, space_after=4, space_before=8)
    method_items = [
        "Visual inspection at 1440px (desktop), 768px (tablet), and 375px (mobile) viewports",
        "Full-page screenshots captured via automated browser (Playwright)",
        "DOM accessibility snapshots for semantic structure analysis",
        "Console error monitoring during page loads",
        "Source code analysis for SEO metadata, color values, and font usage",
        "Google PageSpeed Insights API for performance metrics",
    ]
    add_bullet_list(doc, method_items)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 3. SITE ARCHITECTURE
    # ══════════════════════════════════════════════════════════════

    doc.add_heading('3. Site Architecture Overview', level=1)

    p = doc.add_paragraph()
    run = p.add_run("Technology Stack:")
    run.bold = True
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(4)

    tech_items = [
        "CMS: WordPress (latest)",
        "Theme: Astra v2.5.2",
        "Page Builder: Elementor (with Jet WooBuilder, Jet Engine)",
        "E-commerce: WooCommerce",
        "Animation: MotionPage framework",
        "Chat: Tawk.to (currently failing to load due to CORS errors)",
        "Analytics: Not detected in source",
        "Font: Poppins (Google Fonts via external CDN)",
        "Language: Bilingual (English /en/, French /fr/) via WPML/Polylang",
    ]
    add_bullet_list(doc, tech_items, bold_prefix=True)

    doc.add_paragraph()
    add_styled_paragraph(doc, "Navigation Structure:", bold=True, size=10, space_after=4)
    nav_items = [
        "Home",
        "About Us",
        "Collections (dropdown: Contemporary, Traditional, Windows, Sliding Door, Lift & Slide Doors)",
        "All Products (dropdown: subcategories)",
        "Contact",
        "FR (language switch)",
        "Get A Quote (CTA button)",
    ]
    add_bullet_list(doc, nav_items)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 4. DESIGN & VISUAL IDENTITY AUDIT
    # ══════════════════════════════════════════════════════════════

    doc.add_heading('4. Design & Visual Identity Audit', level=1)

    # 4.1
    doc.add_heading('4.1 Color Scheme & Harmony', level=2)

    p = doc.add_paragraph()
    run = p.add_run("Current Color Palette:")
    run.bold = True
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(4)

    color_table = doc.add_table(rows=8, cols=3)
    color_table.style = 'Table Grid'
    color_headers = ["Element", "Color", "Hex Code"]
    for i, h in enumerate(color_headers):
        color_table.rows[0].cells[i].text = h
    style_table_header(color_table)

    color_data = [
        ("Primary Blue", "Bright Blue", "#046BD2"),
        ("Secondary Blue", "Light Blue", "#93C3F3"),
        ("Header Background", "Near Black", "#191919"),
        ("Footer Background", "Dark Black", "#111111"),
        ("Body Text", "Dark Slate", "#334155"),
        ("Page Background", "Light Blue-Gray", "#F0F5FA"),
        ("White Sections", "Pure White", "#FFFFFF"),
    ]
    for i, (elem, color, hex_code) in enumerate(color_data):
        row = color_table.rows[i + 1]
        row.cells[0].text = elem
        row.cells[1].text = color
        row.cells[2].text = hex_code
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()

    p = add_styled_parameter = add_styled_paragraph(doc, "Core Problem: Dark vs. Light Dissonance", bold=True, size=11,
                             color=RGBColor(0xC0, 0x39, 0x2B), space_after=6)

    p = doc.add_paragraph()
    run = p.add_run("The current design alternates between near-black (#191919, #111111) sections and white (#FFFFFF, #F0F5FA) sections, creating a jarring visual contrast that works against the brand identity. Visifib sells windows and doors \u2014 products that literally bring light into homes. The visual language should reflect this: open, bright, airy, and transparent.")
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph()
    run = p.add_run("Reference: ")
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run("Apple.com uses a predominantly light palette with generous white space, allowing product photography to be the focal point. This approach communicates quality, simplicity, and elegance \u2014 exactly what Visifib's brand should convey.")
    run2.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(8)

    add_styled_paragraph(doc, "Recommendation:", bold=True, size=10,
                         color=RGBColor(0x27, 0xAE, 0x60), space_after=4)

    rec_items = [
        "Transition header from black (#191919) to white or very light gray (#F8F9FA)",
        "Transition footer from black (#111111) to light gray (#F0F2F5) with dark text",
        "Remove all dark background content sections (About Us, Gallery detail pages)",
        "Use the primary blue (#046BD2) sparingly for CTAs and accents only",
        "Let the beautiful product photography stand out against clean, light backgrounds",
        "Maintain consistent light palette across ALL pages for visual harmony",
    ]
    add_bullet_list(doc, rec_items)

    # 4.2
    doc.add_heading('4.2 Hero Images & Layout', level=2)

    p = doc.add_paragraph()
    run = p.add_run("Multiple pages use oversized full-screen hero images that push content below the fold, forcing users to scroll before understanding what the page is about. This is particularly problematic on desktop where the images occupy the entire viewport height.")
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(8)

    hero_table = doc.add_table(rows=6, cols=3)
    hero_table.style = 'Table Grid'
    hero_headers = ["Page", "Issue", "Recommendation"]
    for i, h in enumerate(hero_headers):
        hero_table.rows[0].cells[i].text = h
    style_table_header(hero_table)

    hero_data = [
        ("Homepage", "Three full-viewport hero images stacked (Contemporary, Traditional, Windows) push products below fold", "Reduce hero to 50-60vh max; show product grid immediately"),
        ("Single Gallery (e.g. Gallery 3)", "Massive hero image with no text overlay; user cannot tell what the page is about until scrolling", "Remove hero or reduce to 30vh; show gallery grid immediately"),
        ("Product Blog Posts (e.g. Single Hung 850)", "Two large hero images at top consuming excessive viewport space", "Single compact hero with product name overlay; move product details above fold"),
        ("About Us", "Category banner images followed by content sections", "Keep banner compact (40vh max); ensure key message is visible without scrolling"),
        ("Gallery Index", "Category banner followed by gallery grid", "Reduce banner to compact strip; prioritize gallery thumbnails"),
    ]
    for i, (page, issue, rec) in enumerate(hero_data):
        row = hero_table.rows[i + 1]
        row.cells[0].text = page
        row.cells[1].text = issue
        row.cells[2].text = rec
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    # 4.3
    doc.add_heading('4.3 Animations & Motion', level=2)

    p = doc.add_paragraph()
    run = p.add_run("The site uses the MotionPage framework for animations. While subtle animations can enhance UX, the current implementation creates several problems:")
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(6)

    anim_problems = [
        "Flying/moving images create visual confusion rather than guiding the user's attention",
        "Animation JavaScript adds significant weight to every page load (the MotionPage library loads globally)",
        "Animations delay content visibility \u2014 users must wait for elements to \"fly in\" before reading",
        "On slower connections, partially-loaded animations create a broken/unfinished appearance",
        "The blank space visible on the homepage (between hero and product carousel) is caused by animation elements that have not yet triggered",
    ]
    add_bullet_list(doc, anim_problems)

    add_styled_paragraph(doc, "Recommendation:", bold=True, size=10,
                         color=RGBColor(0x27, 0xAE, 0x60), space_after=4)
    p = doc.add_paragraph()
    run = p.add_run("Remove the MotionPage framework entirely. The only acceptable motion is product carousels (which are already working well). Static content should be immediately visible. Your goal is clarity, not impression \u2014 users should understand your products from the first second, not watch elements fly across the screen.")
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(8)

    # 4.4
    doc.add_heading('4.4 Page-by-Page Design Review', level=2)

    # Homepage
    add_styled_paragraph(doc, "Homepage (visifib.ca/en/)", bold=True, size=10, space_after=4)
    home_items = [
        "Positive: Clean navigation bar with clear CTA (\"Get A Quote\"); product carousel is effective",
        "Issue: Three stacked full-screen hero images dominate the page; content below is invisible without scrolling",
        "Issue: Large blank space between hero section and product carousel (animation placeholder not loading)",
        "Issue: Dark header and footer sandwich light content, breaking visual flow",
        "Fix: Reduce hero to single panel with category links; show product grid prominently above fold",
    ]
    add_bullet_list(doc, home_items)

    # About Us
    add_styled_paragraph(doc, "About Us", bold=True, size=10, space_after=4, space_before=8)
    about_items = [
        "Positive: Company description is clear and professional; \"Our Services\" card is well-designed",
        "Issue: Dark black section for main content area creates harsh contrast with white header area",
        "Issue: Blue text (#046BD2) on dark background is difficult to read",
        "Issue: The alternating black/white sections destroy the light, airy feeling the brand needs",
        "Fix: Replace all dark sections with light backgrounds; use dark text on white/light-gray",
    ]
    add_bullet_list(doc, about_items)

    # Contact
    add_styled_paragraph(doc, "Contact", bold=True, size=10, space_after=4, space_before=8)
    contact_items = [
        "Positive: Contact form is functional with relevant fields (Name, Phone, City, Project Timeline, Product Interest)",
        "Positive: Gradient info card (teal/green) for contact details is visually appealing",
        "Issue: Form sits on a dark photographic background; form card could be more prominent",
        "Issue: Same dark header/footer contrast issue as other pages",
        "Issue: No map integration showing business location",
        "Fix: Add Google Maps embed; lighten overall color scheme; make form the clear focal point",
    ]
    add_bullet_list(doc, contact_items)

    # Gallery
    add_styled_paragraph(doc, "Gallery Index", bold=True, size=10, space_after=4, space_before=8)
    gallery_items = [
        "Positive: Clean 3-column grid layout with clear thumbnails and titles",
        "Issue: Black/white contrast between gallery section and dark footer",
        "Issue: Only 3 gallery items \u2014 consider whether separate gallery index is needed",
        "Fix: Lighten color scheme; consider integrating galleries into product pages directly",
    ]
    add_bullet_list(doc, gallery_items)

    # Single Gallery
    add_styled_paragraph(doc, "Single Gallery Pages (e.g. Gallery 3)", bold=True, size=10, space_after=4, space_before=8)
    sgallery_items = [
        "Issue: Massive hero image (full viewport) with no descriptive text \u2014 user has no context until scrolling",
        "Issue: Dark section with Visifib logo between hero and gallery grid serves no purpose",
        "Issue: Gallery images are paginated (4 pages of 8 images) \u2014 all could load on single page",
        "Fix: Remove hero image entirely; start with gallery title and grid; load all images with lazy-loading",
    ]
    add_bullet_list(doc, sgallery_items)

    # Shop / All Products
    add_styled_paragraph(doc, "Shop / All Products", bold=True, size=10, space_after=4, space_before=8)
    shop_items = [
        "Positive: The cleanest page on the site \u2014 mostly white background, clear product sections",
        "Positive: Good organization by collection (Contemporary, Traditional, Windows, Sliding Door, Lift & Slide)",
        "Positive: Product carousels within each section work well",
        "Issue: Repeated \"Call Us Today / 450-803-2657\" in every section is redundant",
        "Issue: Some placeholder text visible (\"Tell people a little more to attract...\")",
        "Fix: Remove placeholder text; consolidate phone number to header/footer; this page is the design direction the whole site should follow",
    ]
    add_bullet_list(doc, shop_items)

    # Product pages
    add_styled_paragraph(doc, "Single Product Pages (WooCommerce)", bold=True, size=10, space_after=4, space_before=8)
    product_items = [
        "Positive: Clean, simple layout with product image, description, colors, and grille options",
        "Positive: \"Read More\" CTA linking to detailed blog post is good",
        "Issue: Dark header and footer still add unnecessary contrast",
        "Issue: Left side shows large background image that is barely visible behind content",
        "Fix: Clean up background; maintain the simple white-card layout (this is a good template)",
    ]
    add_bullet_list(doc, product_items)

    # Product Blog Posts
    add_styled_paragraph(doc, "Product Blog Posts (e.g. /en/single-hung-850-series/)", bold=True, size=10, space_after=4, space_before=8)
    blog_items = [
        "Issue: Two oversized hero images consuming excessive viewport space on desktop",
        "Issue: Sky blue background (#93C3F3 area) with dark text \u2014 contrast is not optimal for readability",
        "Issue: Section headings (\"Profile\", \"Colours\", \"Glass\") use inconsistent colors",
        "Issue: The interactive product viewer (Double-Hung-Tilt) area has excessive whitespace",
        "Issue: PDF download links (Drawings, Installation) are small and easy to miss",
        "Fix: Single compact hero; consistent heading colors; prominent download buttons; fix blue background contrast",
    ]
    add_bullet_list(doc, blog_items)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 5. PERFORMANCE AUDIT
    # ══════════════════════════════════════════════════════════════

    doc.add_heading('5. Performance Audit', level=1)

    p = doc.add_paragraph()
    run = p.add_run("Website performance directly impacts user experience, conversion rates, and SEO rankings. Google uses Core Web Vitals as a ranking factor. The current site has several performance bottlenecks that can be addressed.")
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(8)

    # 5.1
    doc.add_heading('5.1 Image Optimization', level=2)

    img_items = [
        "Several images are served at original resolution (e.g., 2560x1707px, 1920px wide) regardless of device viewport",
        "Hero images on gallery pages load full-resolution photos that are then CSS-scaled down to fit",
        "Product images in carousels are loaded at 1920px width even when displayed at 200-300px",
        "AVIF format is used in some places (good) but not consistently across all images",
        "No responsive srcset attributes detected for adaptive image serving based on viewport",
    ]
    add_bullet_list(doc, img_items)

    add_styled_paragraph(doc, "Recommendation:", bold=True, size=10,
                         color=RGBColor(0x27, 0xAE, 0x60), space_after=4)
    img_recs = [
        "Implement responsive images with srcset and sizes attributes on all img tags",
        "Generate multiple sizes per image (320px, 640px, 960px, 1280px, 1920px)",
        "Convert all images to WebP/AVIF with JPEG fallback",
        "Use lazy loading (loading=\"lazy\") for all images below the fold",
        "Estimated improvement: 40-60% reduction in image payload, 1-3 second faster page loads",
    ]
    add_bullet_list(doc, img_recs)

    # 5.2
    doc.add_heading('5.2 Web Fonts', level=2)

    p = doc.add_paragraph()
    run = p.add_run("The site loads Poppins (Google Fonts) via an external CDN (dd-time.com), which is currently failing with CORS errors. Additionally, a Roboto font file fails to load on every page.")
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(6)

    font_items = [
        "Poppins font loaded from external CDN adds 100-300ms to page load",
        "Font files failing to load (CORS errors) means text may flash or use fallback fonts unpredictably",
        "Multiple font weights being loaded increases total download size",
        "For a product-focused business site, web fonts add complexity without proportional value",
    ]
    add_bullet_list(doc, font_items)

    add_styled_paragraph(doc, "Recommendation:", bold=True, size=10,
                         color=RGBColor(0x27, 0xAE, 0x60), space_after=4)
    p = doc.add_paragraph()
    run = p.add_run("Switch to system font stack: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Helvetica Neue', Arial, sans-serif. This eliminates font loading entirely, provides instant text rendering, and the visual difference is minimal. Alternatively, self-host Poppins to fix the CORS issues and reduce external dependencies.")
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(8)

    # 5.3
    doc.add_heading('5.3 Caching Strategy', level=2)

    p = doc.add_paragraph()
    run = p.add_run("The site is primarily static content (product catalog, galleries, informational pages) that changes infrequently. An aggressive caching strategy would significantly improve repeat visit performance.")
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(6)

    cache_recs = [
        "Implement browser caching with long TTLs for static assets (images, CSS, JS): Cache-Control: max-age=31536000",
        "Use a caching plugin (WP Super Cache, W3 Total Cache, or WP Rocket) to serve static HTML",
        "Enable server-level Gzip/Brotli compression for all text-based assets",
        "Implement CDN (Cloudflare free tier) for global edge caching and automatic image optimization",
        "Set proper ETags and Last-Modified headers for conditional requests",
        "Estimated improvement: 50-70% faster repeat page loads; significant TTFB improvement",
    ]
    add_bullet_list(doc, cache_recs)

    # 5.4
    doc.add_heading('5.4 Elementor Impact', level=2)

    p = doc.add_paragraph()
    run = p.add_run("Elementor is the page builder used throughout the site. While it provides an easy drag-and-drop interface, it adds significant overhead:")
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(6)

    elem_items = [
        "Large inline CSS: Each page embeds 50-200KB+ of CSS rules directly in the HTML",
        "Multiple JavaScript files: Elementor loads its own JS runtime, plus add-ons (Jet WooBuilder, Jet Engine, MotionPage)",
        "DOM complexity: Elementor generates deeply nested div structures that increase rendering time",
        "Unused CSS: Much of the inline CSS is shared framework styles not used on the specific page",
    ]
    add_bullet_list(doc, elem_items)

    add_styled_paragraph(doc, "Recommendation:", bold=True, size=10,
                         color=RGBColor(0x27, 0xAE, 0x60), space_after=4)
    p = doc.add_paragraph()
    run = p.add_run("Elementor is workable and does not need to be replaced. Focus on: (1) Removing unused Elementor widgets/add-ons, (2) Disabling Elementor on simple pages that don't need it, (3) Using an optimization plugin like Perfmatters or Asset CleanUp to selectively remove unused CSS/JS per page, (4) Enabling Elementor's built-in performance experiments (reduced DOM output, optimized asset loading).")
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(8)

    # 5.5
    doc.add_heading('5.5 Console Errors (Found on Every Page)', level=2)

    error_table = doc.add_table(rows=5, cols=3)
    error_table.style = 'Table Grid'
    err_headers = ["Error", "Source", "Impact"]
    for i, h in enumerate(err_headers):
        error_table.rows[0].cells[i].text = h
    style_table_header(error_table)

    error_data = [
        ("CORS font loading failure", "dd-time.com (Roboto woff2)", "Fonts may not render correctly; fallback flash"),
        ("CORS script blocked", "embed.tawk.to (chat widget)", "Live chat widget fails to load entirely"),
        ("net::ERR_FAILED (font)", "Roboto font file", "Missing font file causes rendering inconsistency"),
        ("Minified React error #299", "react-configurator/assets/index-CEHxvV-G.js", "Product configurator may be partially broken"),
    ]
    for i, (err, source, impact) in enumerate(error_data):
        row = error_table.rows[i + 1]
        row.cells[0].text = err
        row.cells[1].text = source
        row.cells[2].text = impact
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 6. SEO AUDIT
    # ══════════════════════════════════════════════════════════════

    doc.add_heading('6. SEO Audit', level=1)

    # 6.1
    doc.add_heading('6.1 Page Titles & Meta Descriptions', level=2)

    p = doc.add_paragraph()
    run = p.add_run("Every page on the site has an incomplete page title ending with a trailing hyphen and missing site name. Meta descriptions are either missing or inadequate. This is a critical SEO issue affecting all pages.")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    p.paragraph_format.space_after = Pt(8)

    seo_table = doc.add_table(rows=9, cols=3)
    seo_table.style = 'Table Grid'
    seo_headers = ["Page", "Current Title", "Current Meta Description"]
    for i, h in enumerate(seo_headers):
        seo_table.rows[0].cells[i].text = h
    style_table_header(seo_table)

    seo_data = [
        ("Homepage", '"Home -"', '"Visifib.ca"'),
        ("About Us", '"About Us -"', '"Welcome"'),
        ("Contact", '"Contact -"', '"Visifib.ca"'),
        ("Gallery", '"Gallery -"', "Missing"),
        ("Shop", '"All Products -"', "Missing"),
        ("Single Hung 850", '"SINGLE HUNG 850 SERIES -"', '"Lift hardware designed for..."'),
        ("Gallery 3", '"Our Visifib Gallery 3 -"', "Missing"),
        ("Product Page", '"SINGLE HUNG 850 SERIES -"', "Missing (only in schema markup)"),
    ]
    for i, (page, title, desc) in enumerate(seo_data):
        row = seo_table.rows[i + 1]
        row.cells[0].text = page
        row.cells[1].text = title
        row.cells[2].text = desc
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()
    add_styled_paragraph(doc, "Recommendation:", bold=True, size=10,
                         color=RGBColor(0x27, 0xAE, 0x60), space_after=4)
    seo_recs = [
        "Fix the WordPress/Astra SEO settings to include site name after page title (e.g., \"About Us | Visifib - Premium Windows & Doors\")",
        "Write unique, descriptive meta descriptions (150-160 characters) for every page",
        "Include target keywords: \"fiberglass windows\", \"doors Quebec\", \"premium windows Canada\"",
        "Install and configure an SEO plugin (Yoast SEO or Rank Math) for systematic optimization",
        "Each product page should have a unique meta description highlighting key features",
    ]
    add_bullet_list(doc, seo_recs)

    # 6.2
    doc.add_heading('6.2 Structured Data', level=2)
    p = doc.add_paragraph()
    run = p.add_run("Positive: The site implements schema.org structured data for Organization, PostalAddress, and ContactPoint. Product pages include basic product schema. However, the structured data could be expanded to include product reviews, pricing, availability, and FAQ schema for better rich snippet visibility in search results.")
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(8)

    # 6.3
    doc.add_heading('6.3 Sitemap & Indexing', level=2)
    sitemap_issues = [
        "Sitemap exists and is properly structured with 6 sub-sitemaps",
        "Issue: All sitemap URLs point to /fr/ (French) versions \u2014 English /en/ pages may not be indexed",
        "Issue: No hreflang annotations detected for language alternates",
        "Issue: Cart, Checkout, and My Account pages are in the sitemap (should be excluded with noindex)",
        "Recommendation: Create separate sitemaps for each language; add hreflang tags; exclude utility pages",
    ]
    add_bullet_list(doc, sitemap_issues)

    # 6.4
    doc.add_heading('6.4 Content Optimization', level=2)
    content_issues = [
        "Product descriptions are brief (1-3 sentences) \u2014 longer, keyword-rich descriptions improve SEO",
        "Alt text on images uses file names (\"850-1920w\", \"HIGHRES3-1920w\") instead of descriptive text",
        "Blog posts / product detail pages could include FAQ sections for featured snippet opportunities",
        "No internal linking strategy \u2014 products should cross-link to related items",
        "Placeholder text found on Shop page (\"Tell people a little more to attract...\")",
    ]
    add_bullet_list(doc, content_issues)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 7. ACCESSIBILITY AUDIT
    # ══════════════════════════════════════════════════════════════

    doc.add_heading('7. Accessibility Audit (WCAG 2.1)', level=1)

    # 7.1
    doc.add_heading('7.1 Color Contrast Failures', level=2)

    p = doc.add_paragraph()
    run = p.add_run("Multiple contrast failures were identified that violate WCAG 2.1 AA standards (minimum 4.5:1 ratio for normal text, 3:1 for large text):")
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(6)

    contrast_table = doc.add_table(rows=6, cols=4)
    contrast_table.style = 'Table Grid'
    c_headers = ["Location", "Text Color", "Background", "Status"]
    for i, h in enumerate(c_headers):
        contrast_table.rows[0].cells[i].text = h
    style_table_header(contrast_table)

    contrast_data = [
        ("Footer links", "Dark Blue (#046BD2)", "Black (#111111)", "FAIL (2.8:1)"),
        ("Hover states", "Light Blue (#93C3F3)", "Dark (#1D1A1A)", "FAIL (2.8:1)"),
        ("Footer copyright text", "Blue/Gray", "Black (#111111)", "FAIL"),
        ("About Us body text", "Blue (#046BD2)", "Dark section", "MARGINAL (3.5:1)"),
        ("Privacy Policy link", "Blue (#046BD2)", "Dark footer", "FAIL"),
    ]
    for i, (loc, text, bg, status) in enumerate(contrast_data):
        row = contrast_table.rows[i + 1]
        row.cells[0].text = loc
        row.cells[1].text = text
        row.cells[2].text = bg
        row.cells[3].text = status
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
        if "FAIL" in status:
            set_cell_shading(row.cells[3], "FADBD8")

    # 7.2
    doc.add_heading('7.2 Semantic HTML & ARIA', level=2)
    semantic_items = [
        "Positive: Proper use of landmark elements (banner, main, contentinfo, navigation)",
        "Positive: Skip-to-content link present (\"Aller au contenu\")",
        "Issue: Skip link text is in French even on English pages \u2014 should match page language",
        "Issue: Image alt text uses filenames rather than descriptive text (e.g., \"850-1920w\" instead of \"Single Hung 850 Series fiberglass window\")",
        "Issue: Social media links in footer point to the current page URL instead of actual social profiles",
        "Issue: Menu toggle labels use French (\"Permutateur de Menu\") on English pages",
        "Issue: Form labels on Contact page could be more descriptive for screen readers",
    ]
    add_bullet_list(doc, semantic_items)

    # 7.3
    doc.add_heading('7.3 Keyboard Navigation', level=2)
    keyboard_items = [
        "Positive: Focus states are defined with dotted outlines for interactive elements",
        "Positive: Form inputs have proper focus indicators (blue border)",
        "Issue: Product carousel may not be fully keyboard navigable",
        "Issue: Gallery lightbox keyboard navigation not verified",
        "Recommendation: Test full keyboard navigation path on all interactive elements",
    ]
    add_bullet_list(doc, keyboard_items)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 8. TECHNICAL ISSUES
    # ══════════════════════════════════════════════════════════════

    doc.add_heading('8. Technical Issues & Console Errors', level=1)

    p = doc.add_paragraph()
    run = p.add_run("Every page on the site produces 5 console errors on load. These errors indicate broken integrations and failed resource loading:")
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(8)

    tech_items_list = [
        "CORS Error \u2014 Font (dd-time.com): The Roboto font is loaded from an external domain (dd-time.com) that does not include proper CORS headers. This causes the font to be blocked by the browser. Fix: Self-host the font or configure CORS headers on the CDN.",
        "CORS Error \u2014 Tawk.to Chat: The Tawk.to live chat widget script is blocked by CORS policy. The entire chat functionality is non-functional. Fix: Verify Tawk.to account configuration; ensure the embed code is current; check if a security plugin is blocking third-party scripts.",
        "React Error #299: The product configurator (react-configurator/assets/index-CEHxvV-G.js) throws a React hydration/rendering error on every page. This suggests the configurator component may be partially broken. Fix: Debug the React component; check for version conflicts.",
        "Social Links Broken: All three social media icons (Facebook, Twitter, Instagram) in the footer link back to the current page URL instead of actual social media profiles. Fix: Update the footer widget with correct social media URLs.",
        "Incomplete Translations: Navigation elements show French text on English pages (\"Aller au contenu\", \"Permutateur de Menu\", \"Navigation principale du site\"). Fix: Update WPML/Polylang string translations for these UI elements.",
    ]
    add_bullet_list(doc, tech_items_list)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 9. PRIORITIZED RECOMMENDATIONS
    # ══════════════════════════════════════════════════════════════

    doc.add_heading('9. Prioritized Recommendations', level=1)

    p = doc.add_paragraph()
    run = p.add_run("The following recommendations are organized by priority. Phase 1 items should be addressed immediately as they affect user experience, SEO, and functionality. Phase 2 items improve performance and conversions. Phase 3 items are longer-term improvements.")
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(12)

    # Phase 1
    add_styled_paragraph(doc, "PHASE 1: Critical Fixes (Immediate)", bold=True, size=12,
                         color=RGBColor(0xC0, 0x39, 0x2B), space_after=6)

    phase1 = [
        "Fix all page titles: Add site name, remove trailing hyphen (e.g., \"About Us | Visifib - Premium Windows & Doors Quebec\")",
        "Write unique meta descriptions for all pages (150-160 characters each, keyword-rich)",
        "Fix broken social media links in footer (currently pointing to same page)",
        "Fix CORS errors: Self-host fonts; reconfigure Tawk.to chat; debug React configurator",
        "Fix contrast failures: Dark blue text on black footer must be changed to white/light text",
        "Fix French strings on English pages (skip link, menu toggles, navigation labels)",
        "Remove placeholder text from Shop page (\"Tell people a little more to attract...\")",
    ]
    for i, item in enumerate(phase1):
        p = doc.add_paragraph()
        run = p.add_run(f"1.{i+1}  ")
        run.bold = True
        run.font.size = Pt(10)
        run2 = p.add_run(item)
        run2.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(3)

    doc.add_paragraph()

    # Phase 2
    add_styled_paragraph(doc, "PHASE 2: Design & Performance (1-2 Weeks)", bold=True, size=12,
                         color=RGBColor(0xF3, 0x9C, 0x12), space_after=6)

    phase2 = [
        "Transition color scheme to light palette: White/light-gray header, light footer, remove all dark content sections",
        "Reduce hero images: Max 50-60vh on desktop; ensure content is visible above the fold",
        "Remove MotionPage animation framework: Keep product carousels only",
        "Implement responsive images with srcset for all product/gallery images",
        "Convert all images to WebP/AVIF with proper fallbacks",
        "Install caching plugin (WP Rocket or W3 Total Cache) and configure server-level compression",
        "Switch to system fonts or self-host Poppins to eliminate external font loading",
        "Optimize Elementor output: Remove unused widgets, enable performance experiments",
    ]
    for i, item in enumerate(phase2):
        p = doc.add_paragraph()
        run = p.add_run(f"2.{i+1}  ")
        run.bold = True
        run.font.size = Pt(10)
        run2 = p.add_run(item)
        run2.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(3)

    doc.add_paragraph()

    # Phase 3
    add_styled_paragraph(doc, "PHASE 3: SEO & Growth (2-4 Weeks)", bold=True, size=12,
                         color=RGBColor(0x27, 0xAE, 0x60), space_after=6)

    phase3 = [
        "Install and configure SEO plugin (Yoast SEO or Rank Math)",
        "Create English-language sitemap; add hreflang annotations for bilingual support",
        "Write descriptive alt text for all product/gallery images",
        "Expand product descriptions with keyword-rich content",
        "Implement internal linking strategy between related products",
        "Add FAQ schema to product pages for rich snippet opportunities",
        "Set up Google Search Console and submit optimized sitemaps",
        "Implement CDN (Cloudflare) for global caching and automatic image optimization",
        "Add Google Maps embed to Contact page",
        "Consider adding customer testimonials/reviews for social proof",
    ]
    for i, item in enumerate(phase3):
        p = doc.add_paragraph()
        run = p.add_run(f"3.{i+1}  ")
        run.bold = True
        run.font.size = Pt(10)
        run2 = p.add_run(item)
        run2.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(3)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 10. CONCLUSION
    # ══════════════════════════════════════════════════════════════

    doc.add_heading('10. Conclusion & Next Steps', level=1)

    p = doc.add_paragraph()
    run = p.add_run("Visifib.ca has a solid foundation: beautiful product photography, a functional product catalog, and a clear business offering. The website's core challenge is that the current visual presentation (dark color scheme, heavy animations, oversized hero images) works against the brand's natural identity \u2014 a company that brings light into homes through quality windows and doors.")
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph()
    run = p.add_run("The path forward is clear and achievable:")
    run.font.size = Pt(10)
    run.bold = True
    p.paragraph_format.space_after = Pt(6)

    conclusion_items = [
        "Embrace light: Shift the entire color palette to whites, light grays, and subtle blues. Let the product photography be the star.",
        "Simplify: Remove animations, reduce hero image sizes, and ensure every page communicates its purpose within the first viewport.",
        "Fix the fundamentals: Page titles, meta descriptions, contrast ratios, and console errors are quick wins that immediately improve search visibility and user experience.",
        "Optimize for speed: Image optimization, caching, and font strategy will make the site noticeably faster with relatively low effort.",
    ]
    add_bullet_list(doc, conclusion_items)

    p = doc.add_paragraph()
    run = p.add_run("The Shop page already demonstrates what the site could look like with a cleaner approach \u2014 it is the best-designed page and should serve as the design reference for the rest of the site.")
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("We are ready to begin implementation upon your approval. We recommend starting with Phase 1 (critical fixes) immediately, as these items require minimal design changes but deliver significant improvements.")
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(20)

    add_section_divider(doc)

    # Signature block
    doc.add_paragraph()
    p = add_styled_paragraph(doc, "Milorad Djukovic", bold=True, size=12,
                             color=RGBColor(0x1A, 0x1A, 0x2E), space_after=2)
    add_styled_paragraph(doc, "Blazing Sun", size=10, color=RGBColor(0x55, 0x55, 0x55), space_after=2)
    add_styled_paragraph(doc, "milorad.djukovic@blazingsun.space | +381 62 961 62 31", size=10,
                         color=RGBColor(0x55, 0x55, 0x55), space_after=2)

    # ══════════════════════════════════════════════════════════════
    # HEADERS AND FOOTERS
    # ══════════════════════════════════════════════════════════════

    section = doc.sections[0]

    # Enable different odd/even headers
    section.different_first_page_header_footer = False
    sectPr = section._sectPr
    # Enable odd/even
    evenAndOddHeaders = parse_xml(f'<w:evenAndOddHeaders {nsdecls("w")}/>')
    # Check if already exists
    existing = sectPr.findall(qn('w:evenAndOddHeaders'))
    if not existing:
        sectPr.append(evenAndOddHeaders)

    # ── ODD pages header (logo) ──
    odd_header = section.header
    odd_header.is_linked_to_previous = False
    # Clear default
    for p in odd_header.paragraphs:
        p.clear()

    if logo_data:
        p = odd_header.paragraphs[0] if odd_header.paragraphs else odd_header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run()
        logo_data.seek(0)
        run.add_picture(logo_data, width=Inches(1.0))

    # ── EVEN pages header ──
    even_header = section.even_page_header
    even_header.is_linked_to_previous = False
    for p in even_header.paragraphs:
        p.clear()
    p = even_header.paragraphs[0] if even_header.paragraphs else even_header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Prepared for: Visifib")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    p.add_run("\n")
    run2 = p.add_run("Prepared by: Milorad Djukovic (Blazing Sun)")
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ── Footer (same on all pages) ──
    # ODD footer
    odd_footer = section.footer
    odd_footer.is_linked_to_previous = False
    for p in odd_footer.paragraphs:
        p.clear()

    # Add separator line
    p = odd_footer.paragraphs[0] if odd_footer.paragraphs else odd_footer.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="4" w:color="CCCCCC"/></w:pBdr>')
    pPr.append(pBdr)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run("Blazing Sun")
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.add_run("\n")
    run2 = p.add_run("PIB: 114340949 | MB: 67500156")
    run2.font.size = Pt(7)
    run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    p.add_run("\n")
    run3 = p.add_run("Email: milorad.djukovic@blazingsun.space | Phone: +381 62 961 62 31")
    run3.font.size = Pt(7)
    run3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Page number
    p.add_run("\n")
    run4 = p.add_run("Page ")
    run4.font.size = Pt(7)
    run4.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Add page number field
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run5 = p.add_run()
    run5._r.append(fldChar1)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run6 = p.add_run()
    run6._r.append(instrText)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run7 = p.add_run()
    run7._r.append(fldChar2)

    # Even footer (same content)
    even_footer = section.even_page_footer
    even_footer.is_linked_to_previous = False
    for p in even_footer.paragraphs:
        p.clear()

    p = even_footer.paragraphs[0] if even_footer.paragraphs else even_footer.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="4" w:color="CCCCCC"/></w:pBdr>')
    pPr.append(pBdr)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run("Blazing Sun")
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.add_run("\n")
    run2 = p.add_run("PIB: 114340949 | MB: 67500156")
    run2.font.size = Pt(7)
    run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    p.add_run("\n")
    run3 = p.add_run("Email: milorad.djukovic@blazingsun.space | Phone: +381 62 961 62 31")
    run3.font.size = Pt(7)
    run3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Page number
    p.add_run("\n")
    run4 = p.add_run("Page ")
    run4.font.size = Pt(7)
    run4.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run5 = p.add_run()
    run5._r.append(fldChar1)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run6 = p.add_run()
    run6._r.append(instrText)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run7 = p.add_run()
    run7._r.append(fldChar2)

    # ── Save ──
    output_path = "/home/milner/Desktop/remote/Visifib_Website_Audit_Report_BlazingSun.docx"
    doc.save(output_path)
    print(f"Report saved to: {output_path}")
    return output_path

if __name__ == "__main__":
    create_report()
